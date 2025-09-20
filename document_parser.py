import io
import re
from typing import Dict, List, Optional
import streamlit as st

# Document parsing libraries
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None
    st.error("PyMuPDF not available. Please install it for PDF parsing.")

try:
    from docx import Document
except ImportError:
    Document = None
    st.error("python-docx not available. Please install it for DOCX parsing.")

try:
    import docx2txt
except ImportError:
    # Fallback to python-docx if docx2txt is not available
    docx2txt = None

class DocumentParser:
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt']
    
    def extract_text_from_file(self, uploaded_file) -> str:
        """Extract text from uploaded file based on file type"""
        try:
            file_extension = uploaded_file.name.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                return self._extract_from_pdf(uploaded_file)
            elif file_extension == 'docx':
                return self._extract_from_docx(uploaded_file)
            elif file_extension == 'txt':
                return self._extract_from_txt(uploaded_file)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        
        except Exception as e:
            raise Exception(f"Error extracting text from file: {str(e)}")
    
    def _extract_from_pdf(self, uploaded_file) -> str:
        """Extract text from PDF file using PyMuPDF"""
        try:
            # Read the uploaded file
            pdf_bytes = uploaded_file.read()
            
            # Open PDF document
            if fitz is None:
                raise Exception("PyMuPDF not available. Please install it for PDF parsing.")
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            text_content = []
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                
                # Extract text from page
                page_text = page.get_text()
                
                if page_text.strip():
                    text_content.append(page_text)
            
            pdf_document.close()
            
            # Join all pages and clean up
            full_text = "\n".join(text_content)
            return self._clean_extracted_text(full_text)
        
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_from_docx(self, uploaded_file) -> str:
        """Extract text from DOCX file"""
        try:
            # First try with docx2txt if available
            if docx2txt is not None:
                try:
                    text = docx2txt.process(uploaded_file)
                    if text and text.strip():
                        return self._clean_extracted_text(text)
                except Exception:
                    pass
            
            # Fallback to python-docx
            if Document is None:
                raise Exception("python-docx not available. Please install it for DOCX parsing.")
            uploaded_file.seek(0)  # Reset file pointer
            document = Document(uploaded_file)
            
            text_content = []
            
            # Extract paragraphs
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n".join(text_content)
            return self._clean_extracted_text(full_text)
        
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_from_txt(self, uploaded_file) -> str:
        """Extract text from TXT file"""
        try:
            # Read text file with proper encoding
            text_content = uploaded_file.read()
            
            # Try to decode as UTF-8, fallback to latin-1
            try:
                if isinstance(text_content, bytes):
                    text = text_content.decode('utf-8')
                else:
                    text = text_content
            except UnicodeDecodeError:
                if isinstance(text_content, bytes):
                    text = text_content.decode('latin-1')
                else:
                    text = str(text_content)
            
            return self._clean_extracted_text(text)
        
        except Exception as e:
            raise Exception(f"Failed to extract text from TXT: {str(e)}")
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Remove common artifacts
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def parse_job_description(self, text: str) -> Dict[str, Any]:
        """Parse job description text to extract structured information"""
        try:
            parsed_data = {
                'title': self._extract_job_title(text),
                'requirements': self._extract_requirements(text),
                'skills': self._extract_skills(text),
                'qualifications': self._extract_qualifications(text),
                'responsibilities': self._extract_responsibilities(text)
            }
            
            return parsed_data
        
        except Exception as e:
            raise Exception(f"Failed to parse job description: {str(e)}")
    
    def parse_resume(self, text: str) -> Dict[str, Any]:
        """Parse resume text to extract structured information"""
        try:
            parsed_data = {
                'contact_info': self._extract_contact_info(text),
                'education': self._extract_education(text),
                'experience': self._extract_experience(text),
                'skills': self._extract_resume_skills(text),
                'certifications': self._extract_certifications(text),
                'projects': self._extract_projects(text)
            }
            
            return parsed_data
        
        except Exception as e:
            raise Exception(f"Failed to parse resume: {str(e)}")
    
    def _extract_job_title(self, text: str) -> str:
        """Extract job title from job description"""
        lines = text.split('\n')[:10]  # Check first 10 lines
        
        # Look for common job title patterns
        title_patterns = [
            r'job title[:\s]+(.+)',
            r'position[:\s]+(.+)',
            r'role[:\s]+(.+)',
            r'^(.+?)\s*(?:position|role|job)?\s*$'
        ]
        
        for line in lines:
            line = line.strip()
            if len(line) > 5 and len(line) < 100:
                for pattern in title_patterns:
                    match = re.search(pattern, line, re.IGNORECASE)
                    if match:
                        return match.group(1).strip()
        
        return ""
    
    def _extract_requirements(self, text: str) -> List[str]:
        """Extract requirements from job description"""
        requirements = []
        
        # Look for requirements section
        req_patterns = [
            r'requirements?[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)',
            r'qualifications?[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)',
            r'must have[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)'
        ]
        
        for pattern in req_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                req_text = match.group(1).strip()
                # Split by bullet points or new lines
                req_items = re.split(r'[•\-\*]\s*|[\n\r]+', req_text)
                for item in req_items:
                    item = item.strip()
                    if len(item) > 10:
                        requirements.append(item)
        
        return requirements[:10]  # Limit to first 10 requirements
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from job description"""
        skills = []
        
        # Common skill keywords
        skill_keywords = [
            'python', 'java', 'javascript', 'react', 'angular', 'vue',
            'sql', 'mysql', 'postgresql', 'mongodb', 'redis',
            'aws', 'azure', 'gcp', 'docker', 'kubernetes',
            'git', 'linux', 'windows', 'api', 'rest',
            'machine learning', 'ai', 'data science', 'analytics'
        ]
        
        text_lower = text.lower()
        for skill in skill_keywords:
            if skill in text_lower:
                skills.append(skill.title())
        
        # Look for skills section
        skill_patterns = [
            r'skills?[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)',
            r'technologies?[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)',
            r'technical skills?[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)'
        ]
        
        for pattern in skill_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                skill_text = match.group(1).strip()
                # Extract individual skills
                skill_items = re.split(r'[,;•\-\*]|\s+and\s+', skill_text)
                for item in skill_items:
                    item = item.strip()
                    if len(item) > 2 and len(item) < 30:
                        skills.append(item)
        
        return list(set(skills))  # Remove duplicates
    
    def _extract_qualifications(self, text: str) -> List[str]:
        """Extract qualifications from job description"""
        qualifications = []
        
        # Look for education/qualification patterns
        qual_patterns = [
            r'(?:bachelor|master|phd|degree|education)[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)',
            r'qualifications?[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)'
        ]
        
        for pattern in qual_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                qual_text = match.group(1).strip()
                qual_items = re.split(r'[•\-\*]\s*|[\n\r]+', qual_text)
                for item in qual_items:
                    item = item.strip()
                    if len(item) > 10:
                        qualifications.append(item)
        
        return qualifications[:5]  # Limit to first 5 qualifications
    
    def _extract_responsibilities(self, text: str) -> List[str]:
        """Extract responsibilities from job description"""
        responsibilities = []
        
        # Look for responsibilities section
        resp_patterns = [
            r'responsibilities?[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)',
            r'duties?[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)',
            r'you will[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)'
        ]
        
        for pattern in resp_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                resp_text = match.group(1).strip()
                resp_items = re.split(r'[•\-\*]\s*|[\n\r]+', resp_text)
                for item in resp_items:
                    item = item.strip()
                    if len(item) > 10:
                        responsibilities.append(item)
        
        return responsibilities[:8]  # Limit to first 8 responsibilities
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information from resume"""
        contact_info = {}
        
        # Email pattern
        email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        if email_match:
            contact_info['email'] = email_match.group()
        
        # Phone pattern
        phone_match = re.search(r'[\+]?[1-9]?[\d\-\(\)\s]{10,15}', text)
        if phone_match:
            contact_info['phone'] = phone_match.group().strip()
        
        return contact_info
    
    def _extract_education(self, text: str) -> List[str]:
        """Extract education information from resume"""
        education = []
        
        # Look for education section
        edu_patterns = [
            r'education[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)',
            r'academic[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)'
        ]
        
        for pattern in edu_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                edu_text = match.group(1).strip()
                edu_items = re.split(r'[\n\r]+', edu_text)
                for item in edu_items:
                    item = item.strip()
                    if len(item) > 10:
                        education.append(item)
        
        return education[:5]  # Limit to first 5 education entries
    
    def _extract_experience(self, text: str) -> List[str]:
        """Extract work experience from resume"""
        experience = []
        
        # Look for experience section
        exp_patterns = [
            r'experience[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)',
            r'work history[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)',
            r'employment[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)'
        ]
        
        for pattern in exp_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                exp_text = match.group(1).strip()
                exp_items = re.split(r'[\n\r]+', exp_text)
                for item in exp_items:
                    item = item.strip()
                    if len(item) > 15:
                        experience.append(item)
        
        return experience[:8]  # Limit to first 8 experience entries
    
    def _extract_resume_skills(self, text: str) -> List[str]:
        """Extract skills from resume"""
        return self._extract_skills(text)  # Reuse the same logic
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications from resume"""
        certifications = []
        
        # Look for certifications section
        cert_patterns = [
            r'certifications?[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)',
            r'certificates?[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)'
        ]
        
        for pattern in cert_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                cert_text = match.group(1).strip()
                cert_items = re.split(r'[•\-\*]\s*|[\n\r]+', cert_text)
                for item in cert_items:
                    item = item.strip()
                    if len(item) > 5:
                        certifications.append(item)
        
        return certifications[:5]  # Limit to first 5 certifications
    
    def _extract_projects(self, text: str) -> List[str]:
        """Extract projects from resume"""
        projects = []
        
        # Look for projects section
        proj_patterns = [
            r'projects?[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)',
            r'portfolio[:\s]+(.+?)(?=\n\s*[A-Z]|\n\s*\n|\Z)'
        ]
        
        for pattern in proj_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                proj_text = match.group(1).strip()
                proj_items = re.split(r'[•\-\*]\s*|[\n\r]+', proj_text)
                for item in proj_items:
                    item = item.strip()
                    if len(item) > 10:
                        projects.append(item)
        
        return projects[:5]  # Limit to first 5 projects
